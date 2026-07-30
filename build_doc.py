# -*- coding: utf-8 -*-
"""生成 InsWallpaper 图文介绍 Word 文档。"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "doc_assets")
REPO = "https://github.com/njbhvvcvc/InsWallpaper"
COMMIT = "68275760cb69d11f72cd83bca7618c1d25c09b6c"

doc = Document()

# 页面：A4，默认字体微软雅黑
sec = doc.sections[0]
sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(__import__("docx").oxml.ns.qn("w:eastAsia"), "Microsoft YaHei")

def rgb(c):
    if not c:
        return None
    h = c.lstrip("#")
    if len(h) == 3:
        h = "".join(ch * 2 for ch in h)
    return RGBColor.from_string(h)

def hr_title(text, size=24, color="#1F3A56"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = True
    r.font.color.rgb = rgb(color)
    return p

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Microsoft YaHei"
        r.font.color.rgb = rgb("1F3A56")
    return h

def para(text, size=11, bold=False, color=None, align=None, space=4):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = rgb(color)
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(size)
    return p

def pic(name, width=6.3, caption=None):
    path = os.path.join(ASSETS, name)
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = c.add_run(caption); rr.font.size = Pt(9); rr.italic = True
        rr.font.color.rgb = RGBColor.from_string("888888")

# ---------------- 封面 ----------------
hr_title("INS 壁纸轮播器", 30)
hr_title("使用与功能介绍文档", 18, "#5A7CA8")
para("")
para("版本：v3（含底部常驻日志区 · 可自定义双键定格热键）", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color="#555")
para("更新日期：2026-07-29", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color="#555")
para("开源地址：" + REPO, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color="#2F6FB0")
para("")

# ---------------- 一、软件简介（简单） ----------------
heading("一、软件简介", 1)
para("INS 壁纸轮播器是一款运行在 Windows 上的桌面美化小工具。它能把你关注的 "
     "Instagram 账号图片、必应每日壁纸、以及本地文件夹里的照片，自动轮播设为桌面壁纸；"
     "看到喜欢的图，按一下自定义的双键热键即可「定格」固定下来，不用再去翻文件。")
para("它不常驻打扰：平时缩在系统托盘后台运行，并带一个钉在窗口底部的实时日志区，"
     "刷新、抓取、报错都看得见。", color="#444")
para("核心亮点：", bold=True, color="#1F3A56")
bullet("多来源：Instagram 账号 / 必应每日壁纸 / 本地文件夹，可分别设置每日时段。")
bullet("一键定格：默认 F1+F2，可改成任意 ≥2 键组合，随时锁定当前壁纸。")
bullet("零侵入热键：后台线程轮询按键，不影响你在别的软件里正常打字。")
bullet("常驻日志：窗口底部固定黑底日志框，实时滚动，可一键清空。")
bullet("AI 超分：用 Real-ESRGAN 把低分辨率图片放大变清晰。")
bullet("托盘驻留：单实例运行，点 × 收进托盘，右键管理。")

pic("fig1_overview.png", 6.3, "图1  功能总览：八项核心能力一览")

# ---------------- 二、工作流程 ----------------
heading("二、工作流程", 1)
para("从添加图源到一键定格，整体链路如下；运行日志会贯穿全过程，实时反馈状态。")
pic("fig2_workflow.png", 6.3, "图2  工作流程：导入 → 抓取 → 缓存 → 轮播 → 定格，日志全程反馈")

# ---------------- 三、主界面一览 ----------------
heading("三、主界面一览", 1)
para("程序主窗口分为上下两段：上方是可滚动的主内容区（标题 / 天气 / 导入来源 / 来源列表 / "
     "轮播设置 / 控制按钮 / 状态栏），下方是钉在窗口底部的常驻日志区。即使屏幕较矮，"
     "缩窗后日志区也始终可见、不会被裁掉。")
pic("fig3_ui.png", 4.6, "图3  主界面布局示意，橙色高亮处为新增的底部常驻日志区")

# ---------------- 四、详细介绍 ----------------
heading("四、功能详细介绍", 1)

heading("4.1 图片来源", 2)
para("支持三类图源，可同时勾选、分别参与轮播：")
bullet("Instagram 账号：填入账号名，工具用内置无头浏览器抓取该账号的图片。")
bullet("必应每日壁纸：自动拉取 Bing 当天的高清壁纸。")
bullet("本地文件夹：选择一个本地目录，直接轮播其中的图片。")
para("每个来源可单独设置「每日时段」，例如只在晚上 20:00–23:00 轮播某个账号。", color="#444")

heading("4.2 轮播设置", 2)
bullet("轮播间隔：自定义每张壁纸停留的秒数。")
bullet("轮播顺序：顺序播放或随机播放。")
bullet("每来源每日时段：精细化控制不同图源在一天中的展示时间。")

heading("4.3 一键定格热键（重点）", 2)
para("看到喜欢的壁纸，按组合键即可把它「定格」为固定壁纸，暂停轮播。")
bullet("默认 F1 + F2，可在界面里「更换快捷键」改成任意组合。")
bullet("必须 ≥2 个按键同时按下（如 Ctrl+Shift+L、A+B 等），避免误触。")
bullet("采用 GetAsyncKeyState 后台轮询，不挂系统级钩子，因此不会拖慢或卡住其它软件的键盘输入（旧版用键盘钩子曾导致「别的页面打不了字」，已修复）。")

heading("4.4 实时天气（可选）", 2)
para("内置一个可选的天气组件，可显示桌面当前天气并自动更新。它不是强制开启项，"
     "关掉也不影响壁纸轮播。")

heading("4.5 内置音乐播放器", 2)
para("附带一个网页音乐播放器（music_player.html），后台播放、不占用主界面空间，"
     "想听歌时再唤起。")

heading("4.6 运行日志", 2)
para("这是本轮重点改进项：窗口底部固定一条黑底常驻日志框，始终可见。")
bullet("实时滚动：刷新、抓取进度、报错都会实时写入，并自动滚到最新。")
bullet("一键清空：日志框右上角「清空日志」按钮，随时清理。")
bullet("完整查看：按 F12 弹出置顶的完整运行日志 / 报错窗口，便于排查问题。")
bullet("环形缓冲保留最近 1000 条，避免无限占用内存。")

heading("4.7 AI 超分（Real-ESRGAN）", 2)
para("内置 Real-ESRGAN 模型，可对低分辨率图片做超分辨率放大，让老图、小图在 4K "
     "屏幕上更清晰。模型文件随程序一起打包。")

heading("4.8 托盘驻留与单实例", 2)
bullet("单实例：用互斥量 + PID 文件保证同时只跑一个，重复双击会唤醒已有窗口。")
bullet("托盘运行：点窗口 × 默认收进系统托盘后台运行；右键托盘图标可管理、退出。")
bullet("退出兜底：托盘异常时直接安全退出，避免「关不掉只能强删后台」的情况。")

# ---------------- 五、使用方法 ----------------
heading("五、使用方法（快速上手）", 1)
bullet("双击 InsWallpaper.exe 启动（首次会初始化内置浏览器，稍等片刻）。")
bullet("在「导入来源」里添加 Instagram 账号 / 勾选必应每日壁纸 / 添加本地文件夹。")
bullet("在「轮播设置」里设定间隔与顺序，点「开始」即进入自动轮播。")
bullet("看到喜欢的图，按 F1+F2（或你自定义的组合键）一键定格。")
bullet("想看运行状态，直接看窗口底部日志区；按 F12 看完整报错。")
bullet("用完点 × 收进托盘；右键托盘图标可退出。")

# ---------------- 六、技术栈 ----------------
heading("六、技术栈与构建", 1)
bullet("语言/界面：Python + tkinter（GUI）。")
bullet("抓取：Playwright 无头 Chromium（独立子进程，避免主窗口无控制台报错）。")
bullet("热键：GetAsyncKeyState 轮询（零侵入）。")
bullet("超分：Real-ESRGAN（内置模型）。")
bullet("打包：PyInstaller --windowed --onefile，单文件 exe 约 222 MB。")

# ---------------- 七、Git 仓库 ----------------
heading("七、Git 仓库地址", 1)
para("本项目的全部源码与构建脚本已开源托管，欢迎查看、复刻与反馈：")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(REPO); r.font.size = Pt(13); r.bold = True
r.font.color.rgb = RGBColor.from_string("2F6FB0")
para("本次更新已推送至 main 分支，提交号：" + COMMIT, align=WD_ALIGN_PARAGRAPH.CENTER, color="#555")
para("提交链接：" + REPO + "/commit/" + COMMIT, align=WD_ALIGN_PARAGRAPH.CENTER, color="#2F6FB0")

# ---------------- 八、声明 ----------------
heading("八、免责声明", 1)
para("本软件仅供个人将图片设为自己的桌面壁纸使用；图片版权归原作者所有，"
     "请勿用于分发或商业用途。使用 Instagram 图片请遵守对应平台的服务条款。")

out = os.path.join(HERE, "InsWallpaper_介绍文档.docx")
doc.save(out)
print("SAVED", out)
